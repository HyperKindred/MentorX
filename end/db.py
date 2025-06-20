import os
from docx import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LCDocument

# ===== 路径配置 =====
docx_dir = './knowledge'
image_output_dir = './images'
os.makedirs(image_output_dir, exist_ok=True)

all_documents = []

# ===== 遍历所有 docx 文件 =====
for filename in os.listdir(docx_dir):
    if filename.endswith(".docx"):
        file_path = os.path.join(docx_dir, filename)
        doc = Document(file_path)

        # 提取文本
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text_content = '\n'.join(full_text)

        # 保存为 LangChain 文档
        lc_doc = LCDocument(
            page_content=text_content,
            metadata={"source": filename}
        )
        all_documents.append(lc_doc)

        # 提取图片
        for rel in doc.part._rels:
            rel_obj = doc.part._rels[rel]
            if "image" in rel_obj.target_ref:
                image_data = rel_obj.target_part.blob
                image_name = f"{os.path.splitext(filename)[0]}_{rel_obj.target_ref.split('/')[-1]}"
                image_path = os.path.join(image_output_dir, image_name)
                with open(image_path, "wb") as f:
                    f.write(image_data)
                print(f"✅ 提取图片: {image_path}")

# ===== 文本切分 =====
text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
split_docs = text_splitter.split_documents(all_documents)

# ===== 构建向量库 =====
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
db = FAISS.from_documents(split_docs, embedding_model)

# ===== 保存向量库 =====
db.save_local("./multi_doc_vector_db")
print("✅ 向量数据库已保存")
